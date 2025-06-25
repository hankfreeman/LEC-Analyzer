# -*- coding: utf-8 -*-


import os
import sys
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import threading
import queue
import anthropic
import docx
from docx.shared import Pt, RGBColor
import fitz  # PyMuPDF for PDF processing
import time
import re
import logging
import gc
from datetime import datetime
def get_latest_claude_model(client):
    """Get the latest available Claude model"""
    try:
        # Try to get available models from the API
        # Note: This is a fallback approach since Anthropic doesn't expose model listing yet
        # We'll use a list of known models in order of preference (newest first)
        preferred_models = [
            "claude-3-5-sonnet-20241022",  # Latest known as of late 2024
            "self.claude_model",  # Previous version
            "claude-3-sonnet-20240229",    # Fallback
        ]
        
        # Test each model to see which one works
        for model in preferred_models:
            try:
                test_response = client.messages.create(
                    model=model,
                    max_tokens=10,
                    temperature=0,
                    messages=[{"role": "user", "content": "test"}]
                )
                return model  # Return the first working model
            except Exception as e:
                continue
        
        # If none work, return the most conservative fallback
        return "claude-3-sonnet-20240229"
        
    except Exception as e:
        # Default fallback if everything fails
        return "self.claude_model"

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("document_processor.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Processor for Claude API")
        self.root.geometry("800x600")
        self.root.minsize(800, 600)
        
        # API Key handling
        self.api_key = tk.StringVar()
        self.load_api_key()
        
        # Selected files and folders
        self.selected_items = []
        self.processed_results = {}
        self.processing_queue = queue.Queue()
        self.prompts = [
            "Personal History and Living Situation (basic plaintiff facts)",
            "Education History (Plaintiff Only - Do not include Author or Expert Education History",
            "Pre-Event Employment History (Jobs prior to accident)",
            "Employment at the Time of the Event",
            "Post Event Employment (jobs after the accident)",
            "Base Wages (earnings around the time of the accident)",
            "Wage Growth (how much earnings are expected to grow)",
            "Fringe Benefits (Employer-Paid Benefits)",
            "Taxes",
            "Work Life Expectancy (how long plaintiff is expected to work)",
            "Social Security Benefits",
            "Discounting to Present Value (discount rate or inflation)",
            "Loss of Household Services"
        ]
        
        # Model selection - will be determined dynamically
        self.claude_model = None
        
        # Build the UI
        self.setup_ui()

    def clean_claude_response(self, response_text):
        """Clean Claude's response to ensure proper N/A handling
        
        If the response contains 'N/A' anywhere, replace the entire response
        with just 'N/A' to ensure consistency and remove any explanatory text.
        """
        response_text = response_text.strip()
        
        # Check if 'N/A' appears anywhere in the response
        if "N/A" in response_text:
            # If the response isn't exactly 'N/A', replace it
            if response_text != "N/A":
                self.log_progress("Cleaning up response containing N/A")
                return "N/A"
        
        return response_text


    def is_empty_response(self, response_text):
        """Check if the response essentially indicates no information was found"""
        response_text = response_text.lower().strip()
        
        # Patterns that indicate the response is saying "no information found"
        empty_patterns = [
            r"based on my review.*no.*found",
            r"there (is|are|was|were) no",
            r"i (do not|don't|cannot|can't) find",
            r"no relevant information",
            r"no (quotes|information|data|content|text) (related|relevant|pertaining|referring)",
            r"the document does not (contain|include|mention|discuss|address)",
            r"nothing in the document"
        ]
        
        # Check if any pattern matches
        for pattern in empty_patterns:
            if re.search(pattern, response_text):
                return True
        
        return False
    
    def load_api_key(self):
        """Load the API key from a config file or environment variable"""
        try:
            # First try to get from environment
            key = os.environ.get("ANTHROPIC_API_KEY")
            if key:
                self.api_key.set(key)
                return
                
            # Then try from a config file
            if os.path.exists("config.txt"):
                try:
                    with open("config.txt", "r", encoding='utf-8') as f:
                        key = f.read().strip()
                        if key:
                            self.api_key.set(key)
                except UnicodeDecodeError:
                    logger.warning("Config file contains non-UTF-8 characters, trying with different encoding")
                    with open("config.txt", "r", encoding='latin-1') as f:
                        key = f.read().strip()
                        if key:
                            self.api_key.set(key)
        except Exception as e:
            logger.error(f"Error loading API key: {e}")
    
    def save_api_key(self):
        """Save the API key to a config file"""
        try:
            with open("config.txt", "w", encoding='utf-8') as f:
                f.write(self.api_key.get())
            messagebox.showinfo("Success", "API Key saved successfully")
        except Exception as e:
            logger.error(f"Error saving API key: {e}")
            messagebox.showerror("Error", f"Failed to save API key: {e}")
    
    def setup_ui(self):
        """Set up the user interface"""
        # Create a notebook (tabbed interface)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Setup tab
        setup_frame = ttk.Frame(self.notebook)
        self.notebook.add(setup_frame, text="Setup")
        
        # API Key Frame
        api_frame = ttk.LabelFrame(setup_frame, text="Claude API Configuration")
        api_frame.pack(fill=tk.X, expand=False, padx=10, pady=10)
        
        ttk.Label(api_frame, text="API Key:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(api_frame, textvariable=self.api_key, width=50, show="*").grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(api_frame, text="Save API Key", command=self.save_api_key).grid(row=0, column=2, padx=5, pady=5)
        
        # Selection Frame
        selection_frame = ttk.LabelFrame(setup_frame, text="Document Selection - Select Source Documents to Analyze")
        selection_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        button_frame = ttk.Frame(selection_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(button_frame, text="Select Files", command=self.select_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Select Folder", command=self.select_folder).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Clear Selection", command=self.clear_selection).pack(side=tk.LEFT, padx=5)
        
        # List of selected items
        self.selection_listbox = tk.Listbox(selection_frame, selectmode=tk.EXTENDED, height=10)
        self.selection_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Scrollbar for the listbox
        scrollbar = ttk.Scrollbar(self.selection_listbox, orient=tk.VERTICAL, command=self.selection_listbox.yview)
        self.selection_listbox.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Run Frame
        run_frame = ttk.Frame(setup_frame)
        run_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(run_frame, text="Run Processing", command=self.start_processing).pack(side=tk.LEFT, padx=5)
        
        # Progress tab
        progress_frame = ttk.Frame(self.notebook)
        self.notebook.add(progress_frame, text="Progress")
        
        # Progress bar and status
        progress_label_frame = ttk.LabelFrame(progress_frame, text="Processing Progress")
        progress_label_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        ttk.Label(progress_label_frame, text="Current file:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.current_file_label = ttk.Label(progress_label_frame, text="N/A")
        self.current_file_label.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        ttk.Label(progress_label_frame, text="Overall progress:").grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        self.progress_bar = ttk.Progressbar(progress_label_frame, orient=tk.HORIZONTAL, length=300, mode='determinate')
        self.progress_bar.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        
        self.progress_text = tk.Text(progress_label_frame, height=20, width=80, wrap=tk.WORD)
        self.progress_text.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W+tk.E+tk.N+tk.S)
        
        # Configure grid weights to make the text area expandable
        progress_label_frame.grid_rowconfigure(2, weight=1)
        progress_label_frame.grid_columnconfigure(1, weight=1)
        
        # Add scrollbar to the text widget
        text_scrollbar = ttk.Scrollbar(progress_label_frame, orient=tk.VERTICAL, command=self.progress_text.yview)
        self.progress_text.configure(yscrollcommand=text_scrollbar.set)
        text_scrollbar.grid(row=2, column=2, sticky=tk.N+tk.S)
    
    def log_progress(self, message):
        """Log a progress message to both the UI and the log file"""
        logger.info(message)
        self.progress_text.insert(tk.END, f"{message}\n")
        self.progress_text.see(tk.END)
        self.root.update_idletasks()
    
    def select_files(self):
        """Open a file dialog to select multiple documents"""
        files = filedialog.askopenfilenames(
            title="Select documents",
            filetypes=[("Documents", "*.pdf *.docx *.doc *.txt"), ("All files", "*.*")]
        )
        if files:
            for file in files:
                if file not in self.selected_items:
                    self.selected_items.append(file)
                    self.selection_listbox.insert(tk.END, os.path.basename(file))
    
    def select_folder(self):
        """Open a folder dialog to select a folder of documents"""
        folder = filedialog.askdirectory(title="Select folder containing source documents")
        if folder:
            self.selected_items.append(folder)
            self.selection_listbox.insert(tk.END, f"[Folder] {os.path.basename(folder)}")
    
    def clear_selection(self):
        """Clear the current selection"""
        self.selected_items = []
        self.selection_listbox.delete(0, tk.END)
    
    def get_all_document_paths(self):
        """Get all document paths from selected files and folders, avoiding duplicates"""
        document_paths = []
        processed_paths = set()  # Keep track of already processed paths to avoid duplicates
        valid_extensions = ('.pdf', '.docx', '.doc', '.txt')
        
        for item in self.selected_items:
            if os.path.isfile(item) and item.lower().endswith(valid_extensions):
                if os.path.abspath(item) not in processed_paths:
                    document_paths.append(item)
                    processed_paths.add(os.path.abspath(item))
            elif os.path.isdir(item):
                for root, _, files in os.walk(item):
                    for file in files:
                        if file.lower().endswith(valid_extensions):
                            file_path = os.path.join(root, file)
                            if os.path.abspath(file_path) not in processed_paths:
                                document_paths.append(file_path)
                                processed_paths.add(os.path.abspath(file_path))
        
        self.log_progress(f"Found {len(document_paths)} unique documents to process.")
        return document_paths
    
    def find_work_product_folder(self, start_path):
        """Find the Work Product folder by navigating upwards from the start path"""
        current_path = os.path.abspath(start_path)
        
        while True:
            # Check if "Work Product" exists in the current directory
            work_product_path = os.path.join(current_path, "Work Product")
            if os.path.exists(work_product_path) and os.path.isdir(work_product_path):
                return work_product_path
            
            # Move up one directory
            parent_path = os.path.dirname(current_path)
            if parent_path == current_path:  # We've reached the root
                break
            current_path = parent_path
        
        # If not found, create one in the starting directory
        base_dir = os.path.dirname(start_path) if os.path.isfile(start_path) else start_path
        work_product_path = os.path.join(base_dir, "Work Product")
        os.makedirs(work_product_path, exist_ok=True)
        return work_product_path
    
    def extract_text_from_file(self, file_path):
        """Extract text from a file based on its extension with progressive chunking"""
        try:
            if file_path.lower().endswith('.pdf'):
                return self.extract_text_from_pdf(file_path)
            else:
                return {"error": f"Unsupported file format: {os.path.basename(file_path)}. Only PDF files are supported."}
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {"error": f"Error extracting text: {e}"}
    
    def extract_text_from_pdf(self, file_path, max_pages=250):
        """Extract text from a PDF file, with progressive chunking for Expert Reports"""
        try:
            doc = fitz.open(file_path)
            page_count = min(len(doc), max_pages)
            
            # Check if this is an expert report
            is_expert_report = "Expert Reports" in file_path
            
            if is_expert_report:
                # Use progressive chunking for expert reports
                primary_text = ""     # First 10 pages (highest priority)
                secondary_text = ""   # Next 10 pages (medium priority)
                tertiary_text = ""    # Remaining pages (lowest priority)
                
                # Process first 10 pages (primary content)
                for page_num in range(min(10, page_count)):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        # Clean text to handle potential Unicode issues
                        page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                        primary_text += page_text
                        primary_text += f"\n--- Page {page_num + 1} ---\n"
                        logger.info(f"DEBUG: Processed page {page_num + 1} (primary)")
                    except Exception as e:
                        logger.error(f"DEBUG: Error processing page {page_num + 1}: {e}")
                        continue
                
                # Process next 10 pages (secondary content)
                for page_num in range(10, min(20, page_count)):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        # Clean text to handle potential Unicode issues
                        page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                        secondary_text += page_text
                        secondary_text += f"\n--- Page {page_num + 1} ---\n"
                        logger.info(f"DEBUG: Processed page {page_num + 1} (secondary)")
                    except Exception as e:
                        logger.error(f"DEBUG: Error processing page {page_num + 1}: {e}")
                        continue
                
                # Process remaining pages (tertiary content)
                for page_num in range(20, page_count):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        # Clean text to handle potential Unicode issues
                        page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                        tertiary_text += page_text
                        tertiary_text += f"\n--- Page {page_num + 1} ---\n"
                        logger.info(f"DEBUG: Processed page {page_num + 1} (tertiary)")
                    except Exception as e:
                        logger.error(f"DEBUG: Error processing page {page_num + 1}: {e}")
                        continue
                
                return {
                    "is_expert_report": True,
                    "primary_text": primary_text,
                    "secondary_text": secondary_text,
                    "tertiary_text": tertiary_text
                }
            else:
                # Standard approach for non-expert reports
                text = ""
                for page_num in range(page_count):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        # Clean text to handle potential Unicode issues
                        page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                        text += page_text
                        text += f"\n--- Page {page_num + 1} ---\n"
                        logger.info(f"DEBUG: Processed page {page_num + 1}")
                    except Exception as e:
                        logger.error(f"DEBUG: Error processing page {page_num + 1}: {e}")
                        continue
                
                return {
                    "is_expert_report": False,
                    "text": text
                }
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return {"error": f"Error extracting text: {e}"}
        finally:
            # Close the document to free resources
            if 'doc' in locals():
                doc.close()
    
    def reset_claude_session(self, client):
        """Reset the Claude session to prevent any information bleed between documents"""
        # Log the reset
        self.log_progress("Resetting Claude API session between documents")
        
        # Force garbage collection to clean up any memory
        gc.collect()
        
        # Create a clean prompt to reset Claude's context
        reset_prompt = u"""
        This is a new document analysis session. 
        Any information from previous documents or analyses should be disregarded.
        Respond with only 'Session reset confirmed' to acknowledge.
        """
        
        try:
            # Ensure the prompt is properly encoded
            reset_prompt = reset_prompt.encode('utf-8', errors='replace').decode('utf-8')
            
            reset_response = client.messages.create(
                model=self.claude_model,
                max_tokens=100,
                temperature=0,
                system="You are starting a new document analysis session. Any information from previous sessions is irrelevant.",
                messages=[
                    {"role": "user", "content": reset_prompt}
                ]
            )
            
            logger.info("DEBUG: Claude session reset successful")
            return True
        except Exception as e:
            logger.error(f"Error resetting Claude session: {e}")
            self.log_progress(f"Failed to reset Claude session: {e}")
            return False
    
    def process_document_with_claude(self, document_path):
        """Process a document with Claude API, one prompt at a time"""
        self.log_progress(f"Processing: {os.path.basename(document_path)}")
        self.current_file_label.config(text=os.path.basename(document_path))
        
        # Extract text from the document
        document_text_parts = self.extract_text_from_file(document_path)
        
        if isinstance(document_text_parts, dict) and "error" in document_text_parts:
            self.log_progress(f"Failed to extract text from {os.path.basename(document_path)}: {document_text_parts['error']}")
            return None
        
        try:
            # Create the client with the API key
            client = anthropic.Anthropic(api_key=self.api_key.get())
            
            # Determine the best Claude model to use if not already set
            if not self.claude_model:
                self.log_progress("Detecting latest available Claude model...")
                self.claude_model = get_latest_claude_model(client)
                self.log_progress(f"Using Claude model: {self.claude_model}")
            
            # Reset Claude session before processing this document
            self.reset_claude_session(client)
            
            # Process each prompt individually for better results
            results = {}
            
            # Check if this is an expert report with progressive chunking
            is_expert_report = document_text_parts.get("is_expert_report", False)
            
            for prompt in self.prompts:
                self.log_progress(f"Processing prompt '{prompt}' for {os.path.basename(document_path)}")
                
                # For expert reports, process each section separately
                if is_expert_report:
                    # Process primary content first (highest priority)
                    # Ensure all text is properly encoded
                    document_content = document_text_parts.get("primary_text", "")
                    document_content = document_content.encode('utf-8', errors='replace').decode('utf-8')
                    prompt_clean = prompt.encode('utf-8', errors='replace').decode('utf-8')
                    
                    primary_prompt = f"""
                    
                    You are a document analyzer that ONLY provides exact quotes from documents with precise citations.
    
                    RULES:
                    - ONLY return direct quotes from the document. DO NOT explain, summarize, interpret, or provide reasoning.
                    - Do not say "Based on my review", "There was nothing", "The document does not contain", etc.
                    - If NO quotes are found for this topic, respond ONLY with:
                    N/A
                    (Just the letters N/A — nothing else.)
                    
                    EXTRACTION INSTRUCTIONS:
                    - Provide ONLY exact quotes from the document that are directly related to the topic below.
                    - Include quotes that are relevant even if they use different wording.
                    - After each quote, include the page number in this format: (Page X)
                    - If no quotes are found, respond ONLY with "N/A".
                    
                    TOPIC:
                    {prompt_clean}
                    
                    DOCUMENT CONTENT:
                    {document_content}
                    """
                    
                    # Ensure the final prompt is properly encoded
                    primary_prompt = primary_prompt.encode('utf-8', errors='replace').decode('utf-8')
                    
                    self.log_progress(f"Sending primary content for prompt '{prompt}' to Claude API")
                    primary_response = client.messages.create(
                        model=self.claude_model,
                        max_tokens=4000,
                        temperature=0,
                        system="You are a document analyzer that ONLY provides exact quotes from documents with precise citations. CRITICAL INSTRUCTION: If no relevant quotes are found, your ENTIRE response must be ONLY the two characters 'N/A' - no explanations, no reasoning, no other text whatsoever. NEVER explain why you're providing N/A.",
                        messages=[
                            {"role": "user", "content": primary_prompt}
                        ]
                    )
                    primary_result = primary_response.content[0].text
                    
                    # Clean the response
                    if "N/A" in primary_result or self.is_empty_response(primary_result):
                        primary_result = "N/A"
                    
                    # Process secondary content only if primary didn't find enough
                    secondary_result = ""
                    if primary_result.strip() == "N/A" or len(primary_result.strip()) < 100:
                        # Ensure all text is properly encoded
                        document_content = document_text_parts.get("secondary_text", "")
                        document_content = document_content.encode('utf-8', errors='replace').decode('utf-8')
                        
                        secondary_prompt = f"""
                        You are a document analyzer that ONLY provides exact quotes from documents with precise citations.
    
                        RULES:
                        - ONLY return direct quotes from the document. DO NOT explain, summarize, interpret, or provide reasoning.
                        - Do not say "Based on my review", "There was nothing", "The document does not contain", etc.
                        - If NO quotes are found for this topic, respond ONLY with:
                        N/A
                        (Just the letters N/A — nothing else.)
                        
                        EXTRACTION INSTRUCTIONS:
                        - Provide ONLY exact quotes from the document that are directly related to the topic below.
                        - Include quotes that are relevant even if they use different wording.
                        - After each quote, include the page number in this format: (Page X)
                        - If no quotes are found, respond ONLY with "N/A".
                        
                        TOPIC:
                        {prompt_clean}
                        
                        DOCUMENT CONTENT:
                        {document_content}
                        """
                        
                        # Ensure the final prompt is properly encoded
                        secondary_prompt = secondary_prompt.encode('utf-8', errors='replace').decode('utf-8')
                        
                        self.log_progress(f"Sending secondary content for prompt '{prompt}' to Claude API")
                        secondary_response = client.messages.create(
                            model=self.claude_model,
                            max_tokens=4000,
                            temperature=0,
                            system="You are a document analyzer that ONLY provides exact quotes from documents with precise citations. CRITICAL INSTRUCTION: If no relevant quotes are found, your ENTIRE response must be ONLY the two characters 'N/A' - no explanations, no reasoning, no other text whatsoever. NEVER explain why you're providing N/A.",
                            messages=[
                                {"role": "user", "content": secondary_prompt}
                            ]
                        )
                        secondary_result = secondary_response.content[0].text
                        
                        # Clean the response
                        if "N/A" in secondary_result or self.is_empty_response(secondary_result):
                            secondary_result = "N/A"
                        
                        # Wait between requests
                        time.sleep(1)
                    
                    # Process tertiary content only if needed
                    tertiary_result = ""
                    if (primary_result.strip() == "N/A" and secondary_result.strip() == "N/A") or (len(primary_result.strip()) < 100 and len(secondary_result.strip()) < 100 and secondary_result.strip() == "N/A"):
                        # Ensure all text is properly encoded
                        document_content = document_text_parts.get("tertiary_text", "")
                        document_content = document_content.encode('utf-8', errors='replace').decode('utf-8')
                        
                        tertiary_prompt = f"""
                        You are a document analyzer that ONLY provides exact quotes from documents with precise citations.
    
                        RULES:
                        - ONLY return direct quotes from the document. DO NOT explain, summarize, interpret, or provide reasoning.
                        - Do not say "Based on my review", "There was nothing", "The document does not contain", etc.
                        - If NO quotes are found for this topic, respond ONLY with:
                        N/A
                        (Just the letters N/A — nothing else.)
                        
                        EXTRACTION INSTRUCTIONS:
                        - Provide ONLY exact quotes from the document that are directly related to the topic below.
                        - Include quotes that are relevant even if they use different wording.
                        - After each quote, include the page number in this format: (Page X)
                        - If no quotes are found, respond ONLY with "N/A".
                        
                        TOPIC:
                        {prompt_clean}
                        
                        DOCUMENT CONTENT:
                        {document_content}
                        """
                        
                        # Ensure the final prompt is properly encoded
                        tertiary_prompt = tertiary_prompt.encode('utf-8', errors='replace').decode('utf-8')
                        
                        self.log_progress(f"Sending tertiary content for prompt '{prompt}' to Claude API")
                        tertiary_response = client.messages.create(
                            model=self.claude_model,
                            max_tokens=4000,
                            temperature=0,
                            system="You are a document analyzer that ONLY provides exact quotes from documents with precise citations. CRITICAL INSTRUCTION: If no relevant quotes are found, your ENTIRE response must be ONLY the two characters 'N/A' - no explanations, no reasoning, no other text whatsoever. NEVER explain why you're providing N/A.",
                            messages=[
                                {"role": "user", "content": tertiary_prompt}
                            ]
                        )
                        tertiary_result = tertiary_response.content[0].text
                        
                        # Clean the response
                        if "N/A" in tertiary_result or self.is_empty_response(tertiary_result):
                            tertiary_result = "N/A"
                        
                        # Wait between requests
                        time.sleep(1)
                    
                    # Clean up and combine results, properly handling N/A cases
                    all_results = []
                    if primary_result.strip() != "N/A":
                        all_results.append(primary_result.strip())
                    if secondary_result.strip() != "N/A":
                        all_results.append(secondary_result.strip())
                    if tertiary_result.strip() != "N/A":
                        all_results.append(tertiary_result.strip())
                    
                    if all_results:
                        # We have valid content, combine all non-N/A results
                        combined_result = "\n\n".join(all_results)
                    else:
                        # Everything was N/A, just return a single N/A
                        combined_result = "N/A"
                    
                    # Store the combined result
                    results[prompt] = combined_result.strip()
                    self.log_progress(f"Completed processing for prompt '{prompt}'")
                    
                else:
                    # Standard processing for non-expert reports
                    # Ensure all text is properly encoded
                    document_content = document_text_parts.get("text", "")
                    document_content = document_content.encode('utf-8', errors='replace').decode('utf-8')
                    prompt_clean = prompt.encode('utf-8', errors='replace').decode('utf-8')
                    
                    specific_prompt = f"""
                    You are a document analyzer that ONLY provides exact quotes from documents with precise citations.
    
                    RULES:
                    - ONLY return direct quotes from the document. DO NOT explain, summarize, interpret, or provide reasoning.
                    - Do not say "Based on my review", "There was nothing", "The document does not contain", etc.
                    - If NO quotes are found for this topic, respond ONLY with:
                    N/A
                    (Just the letters N/A — nothing else.)
                    
                    EXTRACTION INSTRUCTIONS:
                    - Provide ONLY exact quotes from the document that are directly related to the topic below.
                    - Include quotes that are relevant even if they use different wording.
                    - After each quote, include the page number in this format: (Page X)
                    - If no quotes are found, respond ONLY with "N/A".
                    
                    TOPIC:
                    {prompt_clean}
                    
                    DOCUMENT CONTENT:
                    {document_content}
                    """
                    
                    # Ensure the final prompt is properly encoded
                    specific_prompt = specific_prompt.encode('utf-8', errors='replace').decode('utf-8')
                    
                    self.log_progress(f"Sending prompt '{prompt}' to Claude API")
                    response = client.messages.create(
                        model=self.claude_model,
                        max_tokens=4000,
                        temperature=0,
                        system="You are a document analyzer that ONLY provides exact quotes from documents with precise citations. CRITICAL INSTRUCTION: If no relevant quotes are found, your ENTIRE response must be ONLY the two characters 'N/A' - no explanations, no reasoning, no other text whatsoever. NEVER explain why you're providing N/A.",
                        messages=[
                            {"role": "user", "content": specific_prompt}
                        ]
                    )
                    
                    result = response.content[0].text
                    
                    # Clean the response
                    if "N/A" in result or self.is_empty_response(result):
                        result = "N/A"
                        
                    # Store the result for this prompt
                    results[prompt] = result
                    self.log_progress(f"Received response for prompt '{prompt}'")
                
                # Wait a short time between requests to avoid rate limiting
                time.sleep(1)
            
            # Combine all responses
            combined_response = "\n\n".join([f"## {prompt}\n{response}" for prompt, response in results.items()])
            
            self.log_progress(f"Completed all prompts for: {os.path.basename(document_path)}")
            return {
                "document_name": os.path.basename(document_path),
                "document_path": document_path,
                "response": combined_response,
                "individual_responses": results  # Store individual responses for better processing
            }
        
        except Exception as e:
            self.log_progress(f"Error processing document {os.path.basename(document_path)}: {e}")
            logger.error(f"Error processing document {document_path}: {e}")
            return {
                "document_name": os.path.basename(document_path),
                "document_path": document_path,
                "response": f"Error: {e}"
            }
    
    def parse_claude_response(self, response_data):
        """Parse Claude's response into sections based on prompts without strict validation"""
        parsed_results = {}
        
        # Check if we have individual responses (new method)
        if isinstance(response_data, dict) and "individual_responses" in response_data:
            individual_responses = response_data["individual_responses"]
            for prompt, response in individual_responses.items():
                # Always store the response content regardless of format
                parsed_results[prompt] = {
                    "content": response.strip(),
                    "footnotes": re.findall(r'Page\s+(\d+)', response, re.IGNORECASE)
                }

            return parsed_results
        
        # Fall back to old method if we don't have individual responses
        response = response_data if isinstance(response_data, str) else response_data.get("response", "")
        
        for prompt in self.prompts:
            # Try to find content for this prompt
            pattern = rf"## {re.escape(prompt)}(.+?)(?=\n\n## |$)"
            matches = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
            
            if matches:
                content = matches.group(1).strip()
                
                # Simply check if there's any content
                if content and content != "N/A":
                    parsed_results[prompt] = {
                        "content": content,
                        "footnotes": re.findall(r'Page\s+(\d+)', content, re.IGNORECASE)
                    }
                else:
                    parsed_results[prompt] = {
                        "content": "No information found",
                        "footnotes": []
                    }
            else:
                parsed_results[prompt] = {
                    "content": "No information found",
                    "footnotes": []
                }
        
        return parsed_results
        
    def create_report_document(self, results):
        """Create a Word document with the formatted results without strict validation"""
        try:
            # Create a new document
            doc = docx.Document()
            
            # Add a title
            doc.add_heading('Document Analysis Report', 0)
            
            # Add date and time
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add informational note (not a warning)
            note_para = doc.add_paragraph()
            note_run = note_para.add_run("This report contains information extracted from the analyzed documents.")
            note_run.bold = True
            
            # Add a table of contents heading
            doc.add_heading('Table of Contents', level=1)
            
            # Add table of contents entries
            toc = doc.add_paragraph()
            for prompt in self.prompts:
                # Remove parenthetical information for TOC entries
                clean_prompt = re.sub(r'\s*\([^)]*\)', '', prompt)
                toc.add_run(f"• {clean_prompt}\n")
            
            # Add a page break
            doc.add_page_break()
            
            # Process each prompt
            for prompt in self.prompts:
                # Remove parenthetical information for section headers
                clean_prompt = re.sub(r'\s*\([^)]*\)', '', prompt)
                
                # Add heading for the prompt (without parenthetical info)
                doc.add_heading(clean_prompt, level=1)
                
                has_content = False
                
                # Add information from each document
                for doc_result in results:
                    # Parse the response to extract relevant sections
                    parsed_result = self.parse_claude_response(doc_result)
                    
                    if prompt in parsed_result:
                        content_text = parsed_result[prompt]["content"]
                        # Skip adding document headers for empty or N/A content
                        if content_text and content_text.strip() != "N/A" and content_text != "No information found":
                            has_content = True
                            doc.add_heading(f"From: {doc_result['document_name']}", level=2)
                            p = doc.add_paragraph(content_text)
                            doc.add_paragraph("---")
                
                # If no content was found for this prompt, add a single N/A
                if not has_content:
                    doc.add_paragraph("N/A")
                
                # Add a page break after each section
                doc.add_page_break()
            
            # Find Work Product folder
            if results and results[0]["document_path"]:
                work_product_folder = self.find_work_product_folder(results[0]["document_path"])
            else:
                work_product_folder = self.find_work_product_folder(os.getcwd())
            
            # Save the document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(work_product_folder, f"Document_Analysis_Report_{timestamp}.docx")
            doc.save(output_path)
            
            self.log_progress(f"Report created successfully: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error creating report: {e}")
            self.log_progress(f"Error creating report: {e}")
            return None
    
    def start_processing(self):
        """Start the document processing in a separate thread"""
        if not self.selected_items:
            messagebox.showwarning("Warning", "No documents or folders selected.")
            return
        
        if not self.api_key.get():
            messagebox.showwarning("Warning", "Please enter your Claude API key.")
            return
        
        # Switch to progress tab
        self.notebook.select(1)
        
        # Clear previous progress
        self.progress_text.delete(1.0, tk.END)
        
        # Start processing thread
        threading.Thread(target=self.process_documents, daemon=True).start()
    
    def process_documents(self):
        """Process all documents in the queue with relaxed validation"""
        try:
            document_paths = self.get_all_document_paths()
            
            if not document_paths:
                self.log_progress("No documents selected for processing.")
                return
            
            self.log_progress(f"Processing {len(document_paths)} documents...")
            results = []
            
            # Set up the progress bar
            self.progress_bar["maximum"] = len(document_paths)
            self.progress_bar["value"] = 0
            
            # Process each document without excessive validation
            for i, doc_path in enumerate(document_paths):
                # Process the document
                result = self.process_document_with_claude(doc_path)
                if result:
                    # Always include results if they exist (no validation filters)
                    results.append(result)
                
                # Update progress
                self.progress_bar["value"] = i + 1
                self.root.update_idletasks()
                
                # Brief delay between documents
                time.sleep(1)
            
            if results:
                # Create the report
                self.log_progress("Creating final report...")
                report_path = self.create_report_document(results)
                
                if report_path:
                    self.log_progress("Processing completed successfully!")
                    messagebox.showinfo("Success", f"Processing completed. Report saved to:\n{report_path}")
                else:
                    self.log_progress("Failed to create report.")
                    messagebox.showerror("Error", "Failed to create the report document.")
            else:
                self.log_progress("No results were obtained from document processing.")
                messagebox.showwarning("Warning", "No results were obtained from document processing.")
        
        except Exception as e:
            logger.error(f"Error in document processing: {e}")
            self.log_progress(f"Error in document processing: {e}")
            messagebox.showerror("Error", f"An error occurred during processing: {e}")
        
        finally:
            # Reset UI elements
            self.current_file_label.config(text="N/A")
            self.progress_bar["value"] = 0
            
            # Force garbage collection
            gc.collect()

if __name__ == "__main__":
    # Create the root window
    root = tk.Tk()
    app = DocumentProcessor(root)
    root.mainloop()
